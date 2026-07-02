from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from django.db.models import Avg, Count, ExpressionWrapper, FloatField, F, Sum
from django.utils import timezone
from datetime import timedelta, date
from django.http import HttpResponse
from Uauth.models import Enfant, SuiviEnseignantEnfant
from .models import Lecon, Exercice, EvenementCalendrier
from .serializers import (RechercheEleveSerializer,EleveEnseignantSerializer,LeconCreateSerializer,LeconListSerializer,LeconDetailSerializer,ExerciceSerializer,EvenementCalendrierSerializer,)
from mlt_quiz.models import ScoreQuiz, ThemeQuiz
import io
import re


# VUE : GESTION DES ÉLÈVES (Ajout / Suppression)
# Note : Dépendances d'importation (docx, pypdf, reportlab) installées avec succès.


class RechercheEleveView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    def get(self, request):
        # Vérification du rôle
        if request.user.role != 'ENSEIGNANT':
            return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)
            
        query = request.GET.get('q', '')
        enseignant = request.user.profil_enseignant
        classe_enseignant = enseignant.classe_enseignement
        
        # Filtre les élèves n'appartenant pas encore à sa classe logicielle
        eleves = Enfant.objects.filter(classe=classe_enseignant)
        deja_ajoutes = SuiviEnseignantEnfant.objects.filter(enseignant=enseignant).values_list('enfant_id', flat=True)
        eleves = eleves.exclude(id__in=deja_ajoutes)
        
        if len(query) >= 2:
            from django.db import models as db_models
            eleves = eleves.filter(db_models.Q(utilisateur__first_name__icontains=query) | db_models.Q(utilisateur__last_name__icontains=query) | db_models.Q(utilisateur__username__icontains=query))
        serializer = RechercheEleveSerializer(eleves, many=True)
        return Response(serializer.data)

class EnseignantElevesView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    def get(self, request):
        if request.user.role != 'ENSEIGNANT':
            return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)
        eleves = Enfant.objects.filter(suivienseignantenfant__enseignant__utilisateur=request.user)
        serializer = EleveEnseignantSerializer(eleves, many=True)
        return Response(serializer.data)

    def post(self, request):
        if request.user.role != 'ENSEIGNANT':
            return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)
        eleve_id = request.data.get('eleve_id')
        if not eleve_id:
            return Response({"error": "L'id de l'élève est requis."}, status=status.HTTP_400_BAD_REQUEST)
        try:
            eleve = Enfant.objects.get(id=eleve_id)
        except Enfant.DoesNotExist:
            return Response({"error": "Élève introuvable."}, status=status.HTTP_404_NOT_FOUND)
        enseignant = request.user.profil_enseignant
        if SuiviEnseignantEnfant.objects.filter(enseignant=enseignant, enfant=eleve).exists():
            return Response({"message": "Cet élève est déjà dans votre classe."}, status=status.HTTP_400_BAD_REQUEST)
        SuiviEnseignantEnfant.objects.create(enseignant=enseignant, enfant=eleve)
        return Response({"message": f"{eleve.utilisateur.first_name} {eleve.utilisateur.last_name} a été ajouté à votre classe !"}, status=status.HTTP_201_CREATED)

class SupprimerEleveView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    def delete(self, request, eleve_id):
        if request.user.role != 'ENSEIGNANT':
            return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)
        try:
            suivi = SuiviEnseignantEnfant.objects.get(enseignant__utilisateur=request.user, enfant_id=eleve_id)
        except SuiviEnseignantEnfant.DoesNotExist:
            return Response({"error": "Élève introuvable dans votre classe."}, status=status.HTTP_404_NOT_FOUND)
        suivi.delete()
        return Response({"message": "Élève retiré de votre classe."}, status=status.HTTP_204_NO_CONTENT)

# VUE : STATISTIQUES GLOBAL POUR L'ENSEIGNANT (TABLEAU DE BORD)

class EnseignantStatsView(APIView):
    """
    Donne le résumé de tout ce qui se passe dans la classe : Liste d'élèves, leçons, moyenne de tout le monde.
    """
    permission_classes = [permissions.IsAuthenticated]
    
    def get(self, request):
        if request.user.role != 'ENSEIGNANT':
            return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)
            
        enseignant = request.user.profil_enseignant
        eleves = Enfant.objects.filter(suivienseignantenfant__enseignant=enseignant)
        eleves_ids = eleves.values_list('id', flat=True)
        scores = ScoreQuiz.objects.filter(enfant_id__in=eleves_ids)
        
        # Moyenne globale de toute la classe
        stats_globales = scores.aggregate(moyenne=Avg(ExpressionWrapper(F('points') * 20.0 / F('total_questions'), output_field=FloatField())))
        moyenne_val = round(stats_globales['moyenne'] or 0, 1)
        
        jours_labels = ["Lun", "Mar", "Mer", "Jeu", "Ven", "Sam", "Dim"]
        today = timezone.now().date()
        start_of_week = today - timedelta(days=today.weekday())
        graph_data = []
        
        for i in range(7):
            date_cible = start_of_week + timedelta(days=i)
            entry = {"name": jours_labels[i]}
            for enf in eleves:
                count = scores.filter(enfant=enf, date_realisation=date_cible).count()
                entry[enf.utilisateur.first_name] = count
            graph_data.append(entry)
            
        recent_scores = scores.select_related('enfant__utilisateur').order_by('-date_realisation', '-id')[:5]
        recent_activity = [{"prenom": s.enfant.utilisateur.first_name, "theme": s.get_theme_display(), "score": round((s.points / s.total_questions) * 20, 1), "date": s.date_realisation} for s in recent_scores]
        
        return Response({
            "totalEleves": eleves.count(),
            "totalLecons": Lecon.objects.filter(enseignant=enseignant).count(),
            "totalExercices": Exercice.objects.filter(lecon__enseignant=enseignant).count(),
            "moyenneGenerale": moyenne_val,
            "graphData": graph_data,
            "recentActivity": recent_activity
        }, status=status.HTTP_200_OK)



# VUE : GESTION DES LEÇONS ET EXERCICES

class LeconView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    def get(self, request):
        if request.user.role != 'ENSEIGNANT': return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)
        lecons = Lecon.objects.filter(enseignant__utilisateur=request.user).order_by('-date_creation')
        serializer = LeconListSerializer(lecons, many=True)
        return Response(serializer.data)
        
    def post(self, request):
        if request.user.role != 'ENSEIGNANT': return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)
        serializer = LeconCreateSerializer(data=request.data)
        if serializer.is_valid():
            lecon = serializer.save(enseignant=request.user.profil_enseignant)
            return Response({"message": "Leçon créée avec succès !", "id": lecon.id}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class LeconDetailView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    def _get_lecon(self, request, lecon_id):
        try: return Lecon.objects.get(id=lecon_id, enseignant__utilisateur=request.user)
        except Lecon.DoesNotExist: return None
        
    def get(self, request, lecon_id):
        lecon = self._get_lecon(request, lecon_id)
        if not lecon: return Response({"error": "Leçon introuvable."}, status=status.HTTP_404_NOT_FOUND)
        return Response(LeconDetailSerializer(lecon).data)
        
    def patch(self, request, lecon_id):
        lecon = self._get_lecon(request, lecon_id)
        if not lecon: return Response({"error": "Leçon introuvable."}, status=status.HTTP_404_NOT_FOUND)
        serializer = LeconListSerializer(lecon, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
    def delete(self, request, lecon_id):
        lecon = self._get_lecon(request, lecon_id)
        if not lecon: return Response({"error": "Leçon introuvable."}, status=status.HTTP_404_NOT_FOUND)
        lecon.delete()
        return Response({"message": "Leçon supprimée."}, status=status.HTTP_204_NO_CONTENT)

class ExerciceView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    def _get_lecon(self, request, lecon_id):
        try: return Lecon.objects.get(id=lecon_id, enseignant__utilisateur=request.user)
        except Lecon.DoesNotExist: return None
    def get(self, request, lecon_id):
        lecon = self._get_lecon(request, lecon_id)
        if not lecon: return Response({"error": "Leçon introuvable."}, status=status.HTTP_404_NOT_FOUND)
        exercices = Exercice.objects.filter(lecon=lecon)
        serializer = ExerciceSerializer(exercices, many=True)
        return Response(serializer.data)
    def post(self, request, lecon_id):
        lecon = self._get_lecon(request, lecon_id)
        if not lecon: return Response({"error": "Leçon introuvable."}, status=status.HTTP_404_NOT_FOUND)
        serializer = ExerciceSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save(lecon=lecon)
            return Response({"message": "Exercice ajouté !"}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class ExerciceDeleteView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    def delete(self, request, exercice_id):
        try: exercice = Exercice.objects.get(id=exercice_id, lecon__enseignant__utilisateur=request.user)
        except Exercice.DoesNotExist: return Response({"error": "Exercice introuvable."}, status=status.HTTP_404_NOT_FOUND)
        exercice.delete()
        return Response({"message": "Exercice supprimé."}, status=status.HTTP_204_NO_CONTENT)


# VUE : EXTRACTION DE TEXTE DEPUIS UN PDF OU WORD

class ExtraireTexteView(APIView):
    """
    Reçoit un fichier PDF ou DOCX en multipart/form-data.
    Extrait le texte brut et le retourne en JSON.
    """
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        if request.user.role != 'ENSEIGNANT':
            return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)

        fichier = request.FILES.get('fichier')
        if not fichier:
            return Response({"error": "Aucun fichier fourni."}, status=status.HTTP_400_BAD_REQUEST)

        nom = fichier.name.lower()
        texte = ""

        try:
            if nom.endswith('.pdf'):
                from pypdf import PdfReader
                reader = PdfReader(fichier)
                pages = []
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t.strip())
                texte = "\n\n".join(pages)

            elif nom.endswith('.docx') or nom.endswith('.doc'):
                import docx
                doc = docx.Document(fichier)
                paragraphes = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                texte = "\n\n".join(paragraphes)

            else:
                return Response({"error": "Format non supporté. Utilisez PDF ou DOCX."}, status=status.HTTP_400_BAD_REQUEST)

            if not texte.strip():
                return Response({"error": "Impossible d'extraire du texte de ce fichier."}, status=status.HTTP_422_UNPROCESSABLE_ENTITY)

            # Nettoyer le texte (espaces multiples, lignes vides successives)
            texte = re.sub(r'\n{3,}', '\n\n', texte)
            texte = re.sub(r' {2,}', ' ', texte)

            return Response({"texte": texte, "nb_caracteres": len(texte)}, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({"error": f"Erreur lors de l'extraction : {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# VUE : TÉLÉCHARGEMENT D'UNE LEÇON EN PDF OU WORD

class TelechargementLeconView(APIView):
    """
    Génère et retourne un fichier PDF ou DOCX du contenu d'une leçon avec un formatage premium.
    Paramètre URL : format = 'pdf' ou 'docx' (défaut: pdf)
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, lecon_id):
        if request.user.role != 'ENSEIGNANT':
            return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)

        try:
            lecon = Lecon.objects.get(id=lecon_id, enseignant__utilisateur=request.user)
        except Lecon.DoesNotExist:
            return Response({"error": "Leçon introuvable."}, status=status.HTTP_404_NOT_FOUND)

        format_export = request.GET.get('export_format', 'pdf').lower()
        contenu_brut = lecon.contents if hasattr(lecon, 'contents') else (lecon.contenu or "Aucun contenu disponible.")
        nom_fichier = f"{lecon.titre.replace(' ', '_')}_{lecon.classe}"

        def formater_inline_pdf(texte):
            # Échapper les caractères spéciaux HTML requis par ReportLab
            texte = texte.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            # Restaurer les balises html de formatage après échappement
            # Convertir le gras **gras** -> <b>gras</b>
            texte = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', texte)
            # Convertir l'italique *italique* ou _italique_ -> <i>italique</i>
            texte = re.sub(r'\*([^*]+?)\*', r'<i>\1</i>', texte)
            texte = re.sub(r'_([^_]+?)_', r'<i>\1</i>', texte)
            # Convertir le code inline `code`
            texte = re.sub(r'`([^`]+?)`', r'<font name="Courier" color="#4F46E5"><b>\1</b></font>', texte)
            return texte

        if format_export == 'pdf':
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import cm
                from reportlab.lib import colors
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
                from reportlab.lib.enums import TA_LEFT, TA_CENTER

                buffer = io.BytesIO()
                doc = SimpleDocTemplate(
                    buffer,
                    pagesize=A4,
                    rightMargin=2*cm,
                    leftMargin=2*cm,
                    topMargin=2*cm,
                    bottomMargin=2*cm
                )

                styles = getSampleStyleSheet()
                story = []

                # Style titre principal
                titre_style = ParagraphStyle(
                    'TitreLecon',
                    parent=styles['Title'],
                    fontSize=24,
                    spaceAfter=8,
                    textColor=colors.HexColor('#4F46E5'),
                    fontName='Helvetica-Bold',
                    alignment=TA_LEFT
                )
                
                # Style corps
                corps_style = ParagraphStyle(
                    'Corps',
                    parent=styles['Normal'],
                    fontSize=10.5,
                    leading=18,
                    spaceAfter=8,
                    textColor=colors.HexColor('#475569'),
                    fontName='Helvetica'
                )

                # Style titres de sections (H1, H2, H3)
                h1_style = ParagraphStyle(
                    'H1Style',
                    parent=styles['Heading1'],
                    fontSize=16,
                    spaceBefore=18,
                    spaceAfter=8,
                    textColor=colors.HexColor('#4F46E5'),
                    fontName='Helvetica-Bold'
                )
                h2_style = ParagraphStyle(
                    'H2Style',
                    parent=styles['Heading2'],
                    fontSize=14,
                    spaceBefore=14,
                    spaceAfter=6,
                    textColor=colors.HexColor('#1E293B'),
                    fontName='Helvetica-Bold'
                )
                h3_style = ParagraphStyle(
                    'H3Style',
                    parent=styles['Heading3'],
                    fontSize=12,
                    spaceBefore=10,
                    spaceAfter=4,
                    textColor=colors.HexColor('#475569'),
                    fontName='Helvetica-Bold'
                )

                # Titre de la leçon
                story.append(Paragraph(lecon.titre, titre_style))
                story.append(Spacer(1, 4))

                # Bloc de métadonnées stylisé (Header Banner)
                meta_text = f"<b>Classe :</b> {lecon.classe}  |  <b>Thème :</b> {lecon.theme}  |  <b>Durée :</b> {lecon.duree or '45 min'}"
                meta_style = ParagraphStyle(
                    'MetaStyle',
                    parent=styles['Normal'],
                    fontSize=9.5,
                    textColor=colors.HexColor('#4F46E5'),
                    fontName='Helvetica'
                )
                meta_table = Table([[Paragraph(meta_text, meta_style)]], colWidths=[17*cm])
                meta_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#EEF2FF')),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('LEFTPADDING', (0, 0), (-1, -1), 12),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                    ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#C7D2FE')),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                story.append(meta_table)
                story.append(Spacer(1, 14))

                if lecon.description:
                    desc_style = ParagraphStyle(
                        'DescStyle',
                        parent=corps_style,
                        textColor=colors.HexColor('#6B7280'),
                        fontName='Helvetica-Oblique'
                    )
                    story.append(Paragraph(lecon.description, desc_style))
                    story.append(Spacer(1, 8))

                # Parser le contenu Markdown
                in_blockquote = False
                blockquote_lines = []

                for ligne in contenu_brut.split('\n'):
                    ligne_stripped = ligne.strip()
                    
                    if in_blockquote and not ligne_stripped.startswith('>'):
                        # Générer le callout blockquote accumulé
                        citation_texte = "<br/>".join(blockquote_lines)
                        blockquote_style = ParagraphStyle(
                            'BlockQuoteStyle',
                            parent=styles['Normal'],
                            fontSize=9.5,
                            textColor=colors.HexColor('#5B21B6'),
                            fontName='Helvetica-Oblique',
                            leading=16
                        )
                        quote_table = Table([[Paragraph(citation_texte, blockquote_style)]], colWidths=[17*cm])
                        quote_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F5F3FF')),
                            ('LINELEFT', (0, 0), (0, -1), 4, colors.HexColor('#4F46E5')),
                            ('TOPPADDING', (0, 0), (-1, -1), 10),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                            ('LEFTPADDING', (0, 0), (-1, -1), 14),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 14),
                        ]))
                        story.append(Spacer(1, 8))
                        story.append(quote_table)
                        story.append(Spacer(1, 8))
                        blockquote_lines = []
                        in_blockquote = False

                    if not ligne_stripped:
                        if not in_blockquote:
                            story.append(Spacer(1, 6))
                        else:
                            blockquote_lines.append("")
                    elif ligne_stripped.startswith('>'):
                        in_blockquote = True
                        texte_quote = formater_inline_pdf(ligne_stripped.lstrip('>').strip())
                        blockquote_lines.append(texte_quote)
                    elif ligne_stripped.startswith('### '):
                        story.append(Paragraph(formater_inline_pdf(ligne_stripped[4:]), h3_style))
                    elif ligne_stripped.startswith('## '):
                        story.append(Paragraph(formater_inline_pdf(ligne_stripped[3:]), h2_style))
                    elif ligne_stripped.startswith('# '):
                        story.append(Paragraph(formater_inline_pdf(ligne_stripped[2:]), h1_style))
                    elif ligne_stripped.startswith('- ') or ligne_stripped.startswith('* '):
                        story.append(Paragraph(f"• {formater_inline_pdf(ligne_stripped[2:])}", corps_style))
                    else:
                        story.append(Paragraph(formater_inline_pdf(ligne_stripped), corps_style))

                # Si le fichier finit par une citation
                if in_blockquote and blockquote_lines:
                    citation_texte = "<br/>".join(blockquote_lines)
                    blockquote_style = ParagraphStyle(
                        'BlockQuoteStyleEnd',
                        parent=styles['Normal'],
                        fontSize=9.5,
                        textColor=colors.HexColor('#5B21B6'),
                        fontName='Helvetica-Oblique',
                        leading=16
                    )
                    quote_table = Table([[Paragraph(citation_texte, blockquote_style)]], colWidths=[17*cm])
                    quote_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F5F3FF')),
                        ('LINELEFT', (0, 0), (0, -1), 4, colors.HexColor('#4F46E5')),
                        ('TOPPADDING', (0, 0), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                        ('LEFTPADDING', (0, 0), (-1, -1), 14),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 14),
                    ]))
                    story.append(Spacer(1, 8))
                    story.append(quote_table)

                doc.build(story)
                buffer.seek(0)

                response = HttpResponse(buffer, content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="{nom_fichier}.pdf"'
                return response

            except Exception as e:
                return Response({"error": f"Erreur génération PDF: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        elif format_export == 'docx':
            try:
                import docx
                from docx.shared import Pt, RGBColor, Cm
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls

                doc_word = docx.Document()

                # Marges standardisées
                for sec in doc_word.sections:
                    sec.top_margin = Cm(2.5)
                    sec.bottom_margin = Cm(2.5)
                    sec.left_margin = Cm(2.5)
                    sec.right_margin = Cm(2.5)

                # Titre de la leçon
                p_titre = doc_word.add_paragraph()
                p_titre.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_titre.paragraph_format.space_after = Pt(4)
                run_titre = p_titre.add_run(lecon.titre)
                run_titre.font.name = 'Arial'
                run_titre.font.size = Pt(22)
                run_titre.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                run_titre.bold = True

                # Bloc Métadonnées sous forme de Callout Box (Tableau mono-cellule)
                meta_table = doc_word.add_table(rows=1, cols=1)
                meta_table.autofit = False
                meta_table.columns[0].width = Cm(16)
                meta_cell = meta_table.cell(0, 0)
                
                # Shading & border XML
                meta_shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEF2FF"/>')
                meta_cell._tc.get_or_add_tcPr().append(meta_shading)
                meta_borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="0" w:color="C7D2FE"/><w:top w:val="single" w:sz="12" w:space="0" w:color="C7D2FE"/><w:right w:val="single" w:sz="12" w:space="0" w:color="C7D2FE"/><w:bottom w:val="single" w:sz="12" w:space="0" w:color="C7D2FE"/></w:tcBorders>')
                meta_cell._tc.get_or_add_tcPr().append(meta_borders)

                p_meta = meta_cell.paragraphs[0]
                p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_meta = p_meta.add_run(f"Niveau : {lecon.classe}   |   Thème : {lecon.theme}   |   Durée : {lecon.duree or '45 min'}")
                run_meta.font.name = 'Arial'
                run_meta.font.size = Pt(9.5)
                run_meta.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                run_meta.bold = True

                doc_word.add_paragraph() # Espacement après métadonnées

                if lecon.description:
                    p_desc = doc_word.add_paragraph()
                    p_desc.paragraph_format.space_after = Pt(12)
                    run_desc = p_desc.add_run(lecon.description)
                    run_desc.font.name = 'Arial'
                    run_desc.font.size = Pt(10.5)
                    run_desc.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                    run_desc.italic = True

                def formater_run_docx(p, texte_brut):
                    # Découpage du texte pour appliquer les styles gras, italique et code
                    parties = re.split(r'(\*\*.+?\*\*|\*[^*]+?\*|_[^_]+?_|`[^`]+?`)', texte_brut)
                    for partie in parties:
                        if not partie:
                            continue
                        if partie.startswith('**') and partie.endswith('**'):
                            run = p.add_run(partie[2:-2])
                            run.bold = True
                            run.font.name = 'Arial'
                        elif (partie.startswith('*') and partie.endswith('*')) or (partie.startswith('_') and partie.endswith('_')):
                            run = p.add_run(partie[1:-1])
                            run.italic = True
                            run.font.name = 'Arial'
                        elif partie.startswith('`') and partie.endswith('`'):
                            run = p.add_run(partie[1:-1])
                            run.font.name = 'Consolas'
                            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                            run.bold = True
                        else:
                            run = p.add_run(partie)
                            run.font.name = 'Arial'

                in_blockquote = False
                blockquote_lines = []

                for ligne in contenu_brut.split('\n'):
                    ligne_stripped = ligne.strip()

                    if in_blockquote and not ligne_stripped.startswith('>'):
                        # Création du callout blockquote Word
                        table = doc_word.add_table(rows=1, cols=1)
                        table.autofit = False
                        table.columns[0].width = Cm(16)
                        cell = table.cell(0, 0)
                        
                        # XML shading (fond violet clair) et border (bordure gauche violette)
                        quote_shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F3FF"/>')
                        quote_borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="4F46E5"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
                        cell._tc.get_or_add_tcPr().append(quote_shading)
                        cell._tc.get_or_add_tcPr().append(quote_borders)
                        
                        p_quote = cell.paragraphs[0]
                        p_quote.paragraph_format.left_indent = Cm(0.4)
                        p_quote.paragraph_format.right_indent = Cm(0.4)
                        p_quote.paragraph_format.space_before = Pt(6)
                        p_quote.paragraph_format.space_after = Pt(6)
                        
                        citation_texte = "\n".join(blockquote_lines)
                        run_quote = p_quote.add_run(citation_texte)
                        run_quote.font.name = 'Arial'
                        run_quote.font.size = Pt(10)
                        run_quote.font.color.rgb = RGBColor(0x5B, 0x21, 0xB6)
                        run_quote.italic = True
                        
                        doc_word.add_paragraph()
                        blockquote_lines = []
                        in_blockquote = False

                    if not ligne_stripped:
                        if not in_blockquote:
                            doc_word.add_paragraph()
                        else:
                            blockquote_lines.append("")
                    elif ligne_stripped.startswith('>'):
                        in_blockquote = True
                        blockquote_lines.append(ligne_stripped.lstrip('>').strip())
                    elif ligne_stripped.startswith('### '):
                        texte_titre = ligne_stripped[4:].strip()
                        p = doc_word.add_heading(level=3)
                        p.paragraph_format.space_before = Pt(12)
                        p.paragraph_format.space_after = Pt(4)
                        run = p.runs[0] if p.runs else p.add_run(texte_titre)
                        run.font.name = 'Arial'
                        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                        run.bold = True
                    elif ligne_stripped.startswith('## '):
                        texte_titre = ligne_stripped[3:].strip()
                        p = doc_word.add_heading(level=2)
                        p.paragraph_format.space_before = Pt(16)
                        p.paragraph_format.space_after = Pt(6)
                        run = p.runs[0] if p.runs else p.add_run(texte_titre)
                        run.font.name = 'Arial'
                        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                        run.bold = True
                    elif ligne_stripped.startswith('# '):
                        texte_titre = ligne_stripped[2:].strip()
                        p = doc_word.add_heading(level=1)
                        p.paragraph_format.space_before = Pt(20)
                        p.paragraph_format.space_after = Pt(8)
                        run = p.runs[0] if p.runs else p.add_run(texte_titre)
                        run.font.name = 'Arial'
                        run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                        run.bold = True
                    elif ligne_stripped.startswith('- ') or ligne_stripped.startswith('* '):
                        p = doc_word.add_paragraph(style='List Bullet')
                        p.paragraph_format.space_after = Pt(4)
                        formater_run_docx(p, ligne_stripped[2:])
                    else:
                        p = doc_word.add_paragraph()
                        p.paragraph_format.space_after = Pt(6)
                        formater_run_docx(p, ligne_stripped)

                if in_blockquote and blockquote_lines:
                    table = doc_word.add_table(rows=1, cols=1)
                    table.autofit = False
                    table.columns[0].width = Cm(16)
                    cell = table.cell(0, 0)
                    quote_shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F3FF"/>')
                    quote_borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="4F46E5"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
                    cell._tc.get_or_add_tcPr().append(quote_shading)
                    cell._tc.get_or_add_tcPr().append(quote_borders)
                    
                    p_quote = cell.paragraphs[0]
                    p_quote.paragraph_format.left_indent = Cm(0.4)
                    p_quote.paragraph_format.space_before = Pt(6)
                    p_quote.paragraph_format.space_after = Pt(6)
                    citation_texte = "\n".join(blockquote_lines)
                    run_quote = p_quote.add_run(citation_texte)
                    run_quote.font.name = 'Arial'
                    run_quote.font.size = Pt(10)
                    run_quote.font.color.rgb = RGBColor(0x5B, 0x21, 0xB6)
                    run_quote.italic = True

                buffer = io.BytesIO()
                doc_word.save(buffer)
                buffer.seek(0)

                response = HttpResponse(
                    buffer,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename="{nom_fichier}.docx"'
                return response

            except Exception as e:
                return Response({"error": f"Erreur génération Word: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        else:
            return Response({"error": "Format non supporté. Utilisez 'pdf' ou 'docx'."}, status=status.HTTP_400_BAD_REQUEST)



class EleveDetailView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    def get(self, request, eleve_id):
        if request.user.role != 'ENSEIGNANT': return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)
        try: eleve = Enfant.objects.get(id=eleve_id, suivienseignantenfant__enseignant__utilisateur=request.user)
        except Enfant.DoesNotExist: return Response({"error": "Élève introuvable."}, status=status.HTTP_404_NOT_FOUND)
        serializer = EleveEnseignantSerializer(eleve)
        return Response(serializer.data)

class EleveScoresView(APIView):
    """
    VUE CRITIQUE UNIFIÉE. 
    Retourne exactement la même structure de données statistiques que le composant de parent (`EnfantDetailStatsView`).
    Ceci assure que l'affichage des graphiques est parfait des deux côtés.
    """
    permission_classes = [permissions.IsAuthenticated]
    def get(self, request, eleve_id):
        if request.user.role != 'ENSEIGNANT': return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)
        try:
            # Vérifie que l'élève demandé est bien dans la classe de ce prof
            eleve = Enfant.objects.get(id=eleve_id, suivienseignantenfant__enseignant__utilisateur=request.user)
        except Enfant.DoesNotExist:
            return Response({"error": "Élève introuvable."}, status=status.HTTP_404_NOT_FOUND)

        # 1. Barre de progressions
        stats_themes = ScoreQuiz.objects.filter(enfant=eleve).annotate(
            note_sur_20=ExpressionWrapper(F('points') * 20.0 / F('total_questions'), output_field=FloatField())
        ).values('theme').annotate(moyenne=Avg('note_sur_20'), nb_exercices=Count('id'), temps_moyen=Avg('temps'))

        for stat in stats_themes:
            stat['moyenne'] = round(stat['moyenne'], 2)
            stat['temps_moyen'] = round(stat['temps_moyen'] or 0, 1)
            # CORRECTION DE L'ERREUR 500: Lecture du dictionnaire
            stat['theme_label'] = dict(ThemeQuiz.choices).get(stat['theme'], stat['theme'].capitalize())

        # 2. Construction de l'historique de la page
        scores_query = ScoreQuiz.objects.filter(enfant=eleve).order_by('-date_realisation', '-id')
        historique = [{"id": s.id, "theme": s.get_theme_display(), "note": round((s.points / s.total_questions) * 20, 1), "date": s.date_realisation, "points": s.points, "total": s.total_questions} for s in scores_query]
        
        # 3. Graphique: Ligne (Évolution des notes sur 20)
        progression_notes = [{"date": s.date_realisation.strftime("%d/%m"), "note": round((s.points / s.total_questions) * 20, 1), "theme": s.get_theme_display()} for s in reversed(scores_query[:20])]
        
        # 4. Graphique: Zone (Activité hebdomadaire / Nombre de quiz)
        jours_labels = ["Lun", "Mar", "Mer", "Jeu", "Ven", "Sam", "Dim"]
        today = timezone.now().date()
        start_of_week = today - timedelta(days=today.weekday())
        progression_exercices = []
        for i in range(7):
            date_cible = start_of_week + timedelta(days=i)
            count = scores_query.filter(date_realisation=date_cible).count()
            progression_exercices.append({"name": jours_labels[i], "count": count})

        return Response({
            "enfant": f"{eleve.utilisateur.first_name} {eleve.utilisateur.last_name}",
            "classe": eleve.classe,
            "stats_par_theme": list(stats_themes),
            "historique": historique,
            "progression_notes": progression_notes,
            "progression_exercices": progression_exercices
        }, status=status.HTTP_200_OK)



# VUE : CALENDRIER ENSEIGNANT

class CalendrierView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        if request.user.role != 'ENSEIGNANT':
            return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)

        # ✅ PUBLICATION AUTOMATIQUE PROGRAMMÉE
        # À chaque fois que l'enseignant charge son calendrier, on vérifie
        # si des leçons ont atteint leur date/heure de publication programmée.
        now = timezone.now()
        lecons_a_publier = Lecon.objects.filter(
            enseignant__utilisateur=request.user,
            statut='brouillon',
            date_publication_programmee__lte=now
        )
        if lecons_a_publier.exists():
            for lecon in lecons_a_publier:
                lecon.statut = 'publie'
                lecon.save(update_fields=['statut'])

        # ✅ PUBLICATION VIA CALENDRIER
        # Vérifie aussi les événements marqués "publier_automatiquement"
        # dont la date+heure est passée et la leçon est encore en brouillon
        from datetime import datetime
        evenements_a_publier = EvenementCalendrier.objects.filter(
            enseignant__utilisateur=request.user,
            type_evenement='cours',
            publier_automatiquement=True,
            lecon__statut='brouillon',
            lecon__isnull=False
        )
        for evt in evenements_a_publier:
            if evt.heure:
                dt_evenement = timezone.make_aware(
                    datetime.combine(evt.date, evt.heure),
                    timezone.get_current_timezone()
                )
                if dt_evenement <= now:
                    evt.lecon.statut = 'publie'
                    evt.lecon.save(update_fields=['statut'])

        evenements = EvenementCalendrier.objects.filter(enseignant__utilisateur=request.user)
        serializer = EvenementCalendrierSerializer(evenements, many=True)
        return Response(serializer.data)

    def post(self, request):
        if request.user.role != 'ENSEIGNANT':
            return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)

        data = request.data.copy()

        # Normalisation du titre :
        # Pour un cours, on dérive le titre depuis la leçon si non fourni
        if data.get('type_evenement') == 'cours' and not data.get('titre'):
            lecon_id = data.get('lecon')
            if lecon_id:
                try:
                    lecon_obj = Lecon.objects.get(id=lecon_id, enseignant__utilisateur=request.user)
                    data['titre'] = lecon_obj.titre

                    # Si publication programmée demandée, met à jour la leçon
                    if data.get('publier_automatiquement') and data.get('date') and data.get('heure'):
                        from datetime import datetime
                        try:
                            dt = timezone.make_aware(
                                datetime.strptime(f"{data['date']} {data['heure']}", "%Y-%m-%d %H:%M"),
                                timezone.get_current_timezone()
                            )
                            lecon_obj.date_publication_programmee = dt
                            lecon_obj.save(update_fields=['date_publication_programmee'])
                        except Exception:
                            pass
                except Lecon.DoesNotExist:
                    pass
            else:
                data['titre'] = data.get('titre', 'Événement')

        # Heure vide -> None
        if not data.get('heure'):
            data['heure'] = None

        # lecon_id -> lecon (cas où le frontend envoie l'un ou l'autre)
        if 'lecon_id' in data and 'lecon' not in data:
            data['lecon'] = data.pop('lecon_id')

        serializer = EvenementCalendrierSerializer(data=data)
        if serializer.is_valid():
            evenement = serializer.save(enseignant=request.user.profil_enseignant)
            return Response(EvenementCalendrierSerializer(evenement).data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class CalendrierDeleteView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    def delete(self, request, evenement_id):
        if request.user.role != 'ENSEIGNANT':
            return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)
        try:
            evenement = EvenementCalendrier.objects.get(id=evenement_id, enseignant__utilisateur=request.user)
        except EvenementCalendrier.DoesNotExist:
            return Response({"error": "Événement introuvable."}, status=status.HTTP_404_NOT_FOUND)
        evenement.delete()
        return Response({"message": "Événement supprimé."}, status=status.HTTP_204_NO_CONTENT)


# VUE : PUBLICATION MANUELLE PROGRAMMÉE
# Déclenche immédiatement la vérification et publication des leçons dûment programmées.

class PublierLeconsProgrammeesView(APIView):
    """
    Publie toutes les leçons dont la date_publication_programmee est atteinte.
    Peut être appelée manuellement depuis le frontend ou par un cron.
    """
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        if request.user.role != 'ENSEIGNANT':
            return Response({"error": "Action non autorisée."}, status=status.HTTP_403_FORBIDDEN)

        now = timezone.now()
        lecons = Lecon.objects.filter(
            enseignant__utilisateur=request.user,
            statut='brouillon',
            date_publication_programmee__lte=now
        )
        nb = lecons.count()
        for lecon in lecons:
            lecon.statut = 'publie'
            lecon.save(update_fields=['statut'])
        return Response({
            "message": f"{nb} leçon(s) publiée(s) automatiquement.",
            "publiees": nb
        }, status=status.HTTP_200_OK)


# VUE : EnfantLeconsView

class EnfantLeconsView(APIView):
    """
    Vue pour qu'un enfant puisse voir les leçons
    publiées par son enseignant.
    Filtre : statut = 'publie' + enseignant lié via SuiviEnseignantEnfant.
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):

        if request.user.role != 'ENFANT':
            return Response(
                {"error": "Action non autorisée."},
                status=status.HTTP_403_FORBIDDEN
            )


        # Leçons publiées des enseignants liés à cet enfant
        lecons = Lecon.objects.filter(
            enseignant__eleves__utilisateur=request.user,
            statut='publie'
        ).order_by('-date_creation')

        serializer = LeconListSerializer(lecons, many=True)
        return Response(serializer.data)

# VUE : EnfantLeconDetailView

class EnfantLeconDetailView(APIView):
    """
    Vue pour qu'un enfant puisse voir le détail d'une leçon publiée.
    Vérifie que la leçon appartient bien à son enseignant.
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, lecon_id):

        if request.user.role != 'ENFANT':
            return Response(
                {"error": "Action non autorisée."},
                status=status.HTTP_403_FORBIDDEN
            )

        try:
            lecon = Lecon.objects.get(
                id=lecon_id,
                enseignant__eleves__utilisateur=request.user,
                statut='publie'
            )
        except Lecon.DoesNotExist:
            return Response(
                {"error": "Leçon introuvable."},
                status=status.HTTP_404_NOT_FOUND
            )

        serializer = LeconDetailSerializer(lecon)
        return Response(serializer.data)


# VUE : EnfantExercicesView

class EnfantExercicesView(APIView):
    """
    Vue pour qu'un enfant puisse accéder aux exercices
    d'une leçon publiée par son enseignant.
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, lecon_id):

        if request.user.role != 'ENFANT':
            return Response(
                {"error": "Action non autorisée."},
                status=status.HTTP_403_FORBIDDEN
            )

        try:
            lecon = Lecon.objects.get(
                id=lecon_id,
                enseignant__eleves__utilisateur=request.user,
                statut='publie'
            )
        except Lecon.DoesNotExist:
            return Response(
                {"error": "Leçon introuvable."},
                status=status.HTTP_404_NOT_FOUND
            )

        exercices = Exercice.objects.filter(lecon=lecon)
        serializer = ExerciceSerializer(exercices, many=True)
        return Response(serializer.data)

